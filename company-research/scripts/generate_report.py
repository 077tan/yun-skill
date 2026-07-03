#!/usr/bin/env python3
"""
company-research 深度研究报告生成脚本
根据结构化 JSON 数据生成符合华福证券格式规范的 Word 深度研究报告

用法:
    python generate_report.py --data report_data.json --output 公司名_深度研究报告_日期.docx

JSON 数据结构参见 references/report_schema.md
格式规范参见 references/format_spec.md
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# ─────────────────────────────────────────────
# 颜色常量（经天山铝业报告 XML 实测）
# ─────────────────────────────────────────────
DARK_BLUE   = "1F3864"   # 封面大标题 / 一级标题文字 / 表头背景
MID_BLUE    = "2E75B6"   # 二级标题 / 封面股票代码 / 分隔线
DARK_BLUE2  = "1F4D78"   # 三级标题
GRAY        = "888888"   # 封面英文名 / 辅助文字
TABLE_ALT   = "F2F2F2"   # 表格数据行交替底色（偶行）
TABLE_HEAD  = "1F3864"   # 表头背景 = DARK_BLUE
BORDER_CLR  = "CCCCCC"   # 表格边框色

# ─────────────────────────────────────────────
# 页面常量（DXA: 1 inch = 1440）
# ─────────────────────────────────────────────
PAGE_W      = 12240   # US Letter 宽
PAGE_H      = 15840   # US Letter 高
MARGIN_TB   = 1440    # 上下边距 1 inch
MARGIN_LR   = 1260    # 左右边距 0.875 inch
CONTENT_W   = PAGE_W - MARGIN_LR * 2   # = 9720 DXA
COL_LEFT    = 2400    # 双列表左列宽
COL_RIGHT   = CONTENT_W - COL_LEFT     # = 7320 DXA


# ═══════════════════════════════════════════════
# 低层 XML 工具
# ═══════════════════════════════════════════════

def set_cell_shading(cell, fill_color):
    """设置单元格背景色（CLEAR 模式，避免 SOLID 黑色 bug）"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    # 移除旧的 shd
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    tcPr.append(shd)


def set_cell_borders(cell, color=BORDER_CLR, size=1):
    """设置单元格四边边框（tcBorders 必须在 shd 之前）"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    shd = tcPr.find(qn('w:shd'))
    if shd is not None:
        shd.addprevious(tcBorders)
    else:
        tcPr.append(tcBorders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    """设置单元格内边距（DXA）"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    tcPr.append(tcMar)


def set_cell_width(cell, width_dxa):
    """设置单元格宽度（DXA）"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcPr.insert(0, tcW)


def set_table_width(table, width_dxa=CONTENT_W):
    """强制表格总宽度（tblW 必须是 tblPr 第一个子元素）"""
    tblPr = table._tbl.tblPr
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblPr.insert(0, tblW)


def add_para_border_bottom(para, color=MID_BLUE, size=6, space=4):
    """段落底部加蓝色细线（pBdr 在 pPr 中必须位于 pStyle 之后、spacing 之前）"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), str(space))
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    for existing in pPr.findall(qn('w:pBdr')):
        pPr.remove(existing)
    # pBdr 必须在 spacing/jc/rPr 之前，在 pStyle/numPr 之后
    # 找到第一个不是 pStyle/numPr 的元素，插在它前面
    insert_before = None
    skip_tags = {qn('w:pStyle'), qn('w:numPr')}
    for child in pPr:
        if child.tag not in skip_tags:
            insert_before = child
            break
    if insert_before is not None:
        insert_before.addprevious(pBdr)
    else:
        pPr.append(pBdr)


def add_page_number_field(para):
    """在段落中插入自动页码域"""
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def set_run_font(run, color=None, size_pt=None, bold=False, italic=False, font_name="Arial"):
    """统一设置 run 字体属性"""
    run.font.name = font_name
    run._r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size_pt:
        run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic


def set_spacing(para, before=None, after=None):
    """设置段前段后间距（DXA）- spacing 必须在 jc/ind 之前"""
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        # spacing 应在 jc 之前插入
        jc = pPr.find(qn('w:jc'))
        if jc is not None:
            jc.addprevious(spacing)
        else:
            pPr.append(spacing)
    if before is not None:
        spacing.set(qn('w:before'), str(before))
    if after is not None:
        spacing.set(qn('w:after'), str(after))


# ═══════════════════════════════════════════════
# 页面设置
# ═══════════════════════════════════════════════

def setup_page(doc):
    """设置页面尺寸和边距"""
    section = doc.sections[0]
    section.page_width  = Twips(PAGE_W)
    section.page_height = Twips(PAGE_H)
    section.top_margin    = Twips(MARGIN_TB)
    section.bottom_margin = Twips(MARGIN_TB)
    section.left_margin   = Twips(MARGIN_LR)
    section.right_margin  = Twips(MARGIN_LR)
    section.header_distance = Twips(708)
    section.footer_distance = Twips(708)


# ═══════════════════════════════════════════════
# 页眉 / 页脚
# ═══════════════════════════════════════════════

def build_header(section, company_name, ticker, report_date):
    """
    页眉格式：
    天山铝业集团股份有限公司（002532.SZ）   深度研究报告   2026年3月14日
    ──────────────────────────────── 蓝色细线 ────────────
    """
    header = section.header
    # 清除默认段落
    for p in header.paragraphs:
        p._p.getparent().remove(p._p)

    para = header.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_para_border_bottom(para, color=MID_BLUE, size=6, space=4)

    # 公司名（深蓝）
    r1 = para.add_run(f"{company_name}（{ticker}）")
    set_run_font(r1, color=DARK_BLUE, size_pt=9)

    # 间隔 + 辅助文字（灰色）
    r2 = para.add_run(f"   深度研究报告   {report_date}")
    set_run_font(r2, color=GRAY, size_pt=9)


def build_footer(section):
    """
    页脚格式：居中  第 [页码] 页  |  本报告仅供内部参考，不构成任何投资建议
    """
    footer = section.footer
    for p in footer.paragraphs:
        p._p.getparent().remove(p._p)

    para = footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r1 = para.add_run("第 ")
    set_run_font(r1, color=GRAY, size_pt=9)
    add_page_number_field(para)
    r2 = para.add_run("  |  本报告仅供内部参考，不构成任何投资建议")
    set_run_font(r2, color=GRAY, size_pt=9)


# ═══════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════

def build_cover(doc, meta):
    """
    封面结构：
    （顶部留白）
    公司中文名（28pt 深蓝粗体 居中）
    英文名（14pt 灰色 居中）
    股票代码 | 交易所（12pt 中蓝 居中）
    ─── 蓝色横线 ───
    核心观点：...（12pt 深蓝斜体 居中）
    ─── 蓝色横线 ───
    深度研究报告（14pt 深蓝粗体 居中）
    报告日期（12pt 灰色 居中）
    （底部留白）
    免责声明（9pt 灰色 居中）
    分页符
    """
    # 顶部留白
    p = doc.add_paragraph()
    set_spacing(p, before=1800)

    # 公司中文名
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p1.add_run(meta['company_name_cn'])
    set_run_font(r, color=DARK_BLUE, size_pt=28, bold=True)

    # 英文名
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(meta.get('company_name_en', ''))
    set_run_font(r, color=GRAY, size_pt=14)

    # 股票代码行
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run(f"{meta['ticker']}  |  {meta['exchange']}")
    set_run_font(r, color=MID_BLUE, size_pt=12)

    # 上分隔线（空段落底部加线）
    sep1 = doc.add_paragraph()
    sep1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(sep1, before=500, after=200)
    add_para_border_bottom(sep1, color=MID_BLUE, size=12, space=0)

    # 核心观点
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p4, after=200)
    r = p4.add_run(f"核心观点：{meta['core_view']}")
    set_run_font(r, color=DARK_BLUE, size_pt=12, italic=True)

    # 下分隔线
    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(sep2, after=500)
    _add_top_border(sep2, color=MID_BLUE, size=12)

    # 报告类型
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p5.add_run("深度研究报告")
    set_run_font(r, color=DARK_BLUE, size_pt=14, bold=True)

    # 报告日期
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p6.add_run(f"报告日期：{meta['report_date']}")
    set_run_font(r, color=GRAY, size_pt=12)

    # 底部留白 + 免责声明
    gap = doc.add_paragraph()
    set_spacing(gap, before=1200)
    p7 = doc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p7.add_run("免责声明：本报告仅供内部参考，不构成任何形式的投资建议。")
    set_run_font(r, color=GRAY, size_pt=9)

    # 分页符
    doc.add_page_break()


def _add_top_border(para, color=MID_BLUE, size=12):
    """段落顶部加线（pBdr 位于 pStyle 之后、spacing 之前）"""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), str(size))
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), color)
    pBdr.append(top)
    for existing in pPr.findall(qn('w:pBdr')):
        pPr.remove(existing)
    skip_tags = {qn('w:pStyle'), qn('w:numPr')}
    insert_before = None
    for child in pPr:
        if child.tag not in skip_tags:
            insert_before = child
            break
    if insert_before is not None:
        insert_before.addprevious(pBdr)
    else:
        pPr.append(pBdr)


# ═══════════════════════════════════════════════
# 标题工具
# ═══════════════════════════════════════════════

def add_h1(doc, text):
    """一级标题：16pt 深蓝粗体 + 底部蓝色细线"""
    p = doc.add_paragraph()
    set_spacing(p, before=360, after=120)
    add_para_border_bottom(p, color=MID_BLUE, size=6, space=4)
    r = p.add_run(text)
    set_run_font(r, color=DARK_BLUE, size_pt=16, bold=True)
    return p


def add_h2(doc, text):
    """二级标题：13pt 中蓝粗体"""
    p = doc.add_paragraph()
    set_spacing(p, before=180, after=120)
    r = p.add_run(text)
    set_run_font(r, color=MID_BLUE, size_pt=13, bold=True)
    return p


def add_body(doc, text, size_pt=11):
    """正文段落：11pt Arial，段前3pt，段后3pt"""
    p = doc.add_paragraph()
    set_spacing(p, before=60, after=60)   # 60 DXA = 3pt
    r = p.add_run(text)
    set_run_font(r, size_pt=size_pt)
    return p


def add_label(doc, text):
    """小节内标签行（如"利润表摘要"/"2020年"）：Bold + 深蓝 #1F3864，段前9pt，段后3pt"""
    p = doc.add_paragraph()
    set_spacing(p, before=180, after=60)  # 180 DXA = 9pt, 60 DXA = 3pt
    r = p.add_run(text)
    set_run_font(r, color=DARK_BLUE, bold=True)
    return p


def add_note(doc, text):
    """注释行：9.5pt 灰色，段前2pt，段后2pt"""
    p = doc.add_paragraph()
    set_spacing(p, before=40, after=40)   # 40 DXA ≈ 2pt
    r = p.add_run(text)
    set_run_font(r, color=GRAY, size_pt=9.5)
    return p


def add_bullet(doc, text, size_pt=11):
    """bullet 列表项（手动实现，避免 List Bullet style 的 w:left 问题）"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    # 手动设置缩进（用 w:ind 而非 w:left）
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '360')
    ind.set(qn('w:hanging'), '360')
    pPr.append(ind)
    r = p.add_run("• " + text)
    set_run_font(r, size_pt=size_pt)
    return p


# ═══════════════════════════════════════════════
# 表格工具
# ═══════════════════════════════════════════════

def make_info_table(doc, rows_data):
    """
    双列信息表（项目 | 内容）
    左列 2400 DXA，右列 7320 DXA
    第一行为表头（深蓝底白字）
    """
    table = doc.add_table(rows=0, cols=2)
    set_table_width(table, CONTENT_W)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 设置列宽
    col_widths = [COL_LEFT, COL_RIGHT]
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            set_cell_width(cell, col_widths[i])

    for i, (label, content) in enumerate(rows_data):
        row = table.add_row()
        is_header = (i == 0)
        for j, (cell, val) in enumerate(zip(row.cells, [label, content])):
            set_cell_shading(cell, TABLE_HEAD if is_header else ("FFFFFF" if i % 2 == 1 else TABLE_ALT))
            set_cell_borders(cell)
            set_cell_margins(cell)
            set_cell_width(cell, col_widths[j])
            p = cell.paragraphs[0]
            if j == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            set_run_font(r,
                         color="FFFFFF" if is_header else None,
                         size_pt=10,
                         bold=is_header)
    return table


def make_data_table(doc, headers, rows, col_widths=None):
    """
    多列数据表（财务数据 / 股价表现等）
    headers: list[str]
    rows: list[list]
    col_widths: list[int DXA]，默认均分
    """
    n_cols = len(headers)
    if col_widths is None:
        w = CONTENT_W // n_cols
        col_widths = [w] * n_cols
        col_widths[-1] = CONTENT_W - w * (n_cols - 1)  # 尾列补齐

    table = doc.add_table(rows=0, cols=n_cols)
    set_table_width(table, CONTENT_W)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 表头行
    hdr_row = table.add_row()
    for j, (cell, hdr) in enumerate(zip(hdr_row.cells, headers)):
        set_cell_shading(cell, TABLE_HEAD)
        set_cell_borders(cell)
        set_cell_margins(cell)
        set_cell_width(cell, col_widths[j])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(hdr))
        set_run_font(r, color="FFFFFF", size_pt=10, bold=True)

    # 数据行
    for i, row_data in enumerate(rows):
        row = table.add_row()
        fill = "FFFFFF" if i % 2 == 0 else TABLE_ALT
        for j, (cell, val) in enumerate(zip(row.cells, row_data)):
            set_cell_shading(cell, fill)
            set_cell_borders(cell)
            set_cell_margins(cell)
            set_cell_width(cell, col_widths[j])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val) if val is not None else "—")
            set_run_font(r, size_pt=10)

    return table


# ═══════════════════════════════════════════════
# 各章节构建函数
# ═══════════════════════════════════════════════

def build_section1(doc, data):
    """第一部分：公司基本情况"""
    add_h1(doc, "第一部分：公司基本情况")

    # 1.1 企业基本信息
    add_h2(doc, "1.1 企业基本信息")
    info = data.get('company_info', {})
    info_rows = [
        ("项目", "内容"),
        ("成立时间", info.get('founded', '')),
        ("上市时间", info.get('listed', '')),
        ("股票代码", info.get('ticker', '')),
        ("注册地", info.get('hq', '')),
        ("企业性质", info.get('ownership', '')),
        ("实际控制人", info.get('controller', '')),
        ("主要生产/业务基地", info.get('bases', '')),
        ("近期重大资本运作", info.get('capital_events', '无')),
    ]
    make_info_table(doc, info_rows)
    doc.add_paragraph()  # 间距

    # 所属指数
    if data.get('indices'):
        add_label(doc, "所属重要股票指数：")
        idx_rows = [("指数名称", "备注")] + [(r.get('name', ''), r.get('note', '')) for r in data['indices']]
        make_info_table(doc, idx_rows)
        doc.add_paragraph()

    # 1.2 股价表现
    add_h2(doc, "1.2 股价表现总览")
    perf = data.get('price_performance', [])
    if perf:
        headers = ["年度", f"{data['meta']['company_name_cn']}年度涨跌幅",
                   f"{data.get('benchmark_name', '基准指数')}年度涨跌幅", "相对收益", "主要驱动"]
        col_widths = [900, 1800, 1800, 1500, 3720]
        rows = [[r.get('year'), r.get('stock_return'), r.get('index_return'),
                 r.get('excess_return'), r.get('driver')] for r in perf]
        make_data_table(doc, headers, rows, col_widths)
        doc.add_paragraph()

    # 1.3 财务概况
    add_h2(doc, "1.3 财务概况（压缩报表）")
    fin = data.get('financials', {})
    years = fin.get('years', [])
    if years:
        _build_financials_table(doc, fin, years)
        doc.add_paragraph()

    # 1.4 社会舆论评价
    add_h2(doc, "1.4 社会舆论评价")
    for para_text in data.get('reputation', []):
        add_body(doc, para_text)


def _build_financials_table(doc, fin, years):
    """利润表 + 资产负债表摘要"""
    # 利润表
    add_label(doc, "利润表摘要（单位：亿元）")
    metrics_is = [
        ("营业收入（亿元）", "revenue"),
        ("收入同比增速",     "revenue_yoy"),
        ("毛利（亿元）",     "gross_profit"),
        ("毛利率",           "gross_margin"),
        ("四费费用率",       "opex_ratio"),
        ("营业利润（亿元）", "op_profit"),
        ("营业利润率",       "op_margin"),
        ("归母净利润（亿元）", "net_profit"),
        ("归母净利润增速",   "net_profit_yoy"),
        ("归母净利率",       "net_margin"),
    ]
    year_cols = [int(CONTENT_W * 0.22)] * len(years)
    label_col = CONTENT_W - sum(year_cols)
    col_widths = [label_col] + year_cols

    headers = ["科目"] + [str(y) for y in years]
    rows = []
    is_data = fin.get('income_statement', {})
    for label, key in metrics_is:
        row = [label]
        for y in years:
            row.append(is_data.get(str(y), {}).get(key, "—"))
        rows.append(row)
    make_data_table(doc, headers, rows, col_widths)
    doc.add_paragraph()

    # 资产负债表
    add_label(doc, "资产负债表摘要（单位：亿元）")
    metrics_bs = [
        ("现金+固收+理财",       "cash"),
        ("经营性应收",           "trade_receivables"),
        ("存货",                 "inventory"),
        ("固定资产",             "ppe"),
        ("在建工程",             "construction"),
        ("无形资产",             "intangibles"),
        ("商誉+权益性投资",      "goodwill_investments"),
        ("资产总计",             "total_assets"),
        ("短期负债（短借+到期）", "short_debt"),
        ("长期借款+应付债券",    "long_debt"),
        ("预收/合同负债",        "contract_liabilities"),
        ("经营性应付",           "trade_payables"),
        ("归母所有者权益",       "equity"),
    ]
    rows_bs = []
    bs_data = fin.get('balance_sheet', {})
    for label, key in metrics_bs:
        row = [label]
        for y in years:
            row.append(bs_data.get(str(y), {}).get(key, "—"))
        rows_bs.append(row)
    make_data_table(doc, headers, rows_bs, col_widths)


def build_section2(doc, data):
    """第二部分：年报逐年阅读与总结"""
    add_h1(doc, "第二部分：年报逐年阅读与总结")
    add_h2(doc, "2.1 逐年经营情况")

    for yr in data.get('annual_review', []):
        # 年份小标题：Bold + 深蓝 #1F3864，与样本报告一致
        add_label(doc, yr.get('title', ''))

        # 核心财务数据表（小表）
        kpi = yr.get('kpis', {})
        if kpi:
            add_label(doc, "核心财务数据")
            kpi_rows = [("指标", "数据")] + [(k, v) for k, v in kpi.items()]
            make_info_table(doc, kpi_rows)
            doc.add_paragraph()

        # 正文内容
        for para_text in yr.get('content', []):
            add_body(doc, para_text)

    # 2.2 综合总结
    add_h2(doc, "2.2 综合总结")
    summary = data.get('annual_summary', {})
    for section_title, paras in summary.items():
        add_label(doc, section_title)
        for para_text in paras:
            add_body(doc, para_text)


def build_section3(doc, data):
    """第三部分：股价复盘"""
    add_h1(doc, "第三部分：股价复盘")

    for phase in data.get('price_review', []):
        add_h2(doc, phase.get('title', ''))
        for para_text in phase.get('content', []):
            add_body(doc, para_text)

    # 关键拐点总结表
    pivots = data.get('key_pivots', [])
    if pivots:
        add_h2(doc, "关键拐点总结")
        headers = ["时间点", "价格", "类型", "判断依据与背后逻辑"]
        col_widths = [1200, 900, 900, 6720]
        rows = [[p.get('date'), p.get('price'), p.get('type'), p.get('reason')] for p in pivots]
        make_data_table(doc, headers, rows, col_widths)
        doc.add_paragraph()
        add_note(doc, "注：以上判断均基于历史数据事后分析（推断），不代表彼时信号的实时可识别性。")


def build_section4(doc, data):
    """第四部分：定性认识"""
    add_h1(doc, "第四部分：思考第一阶段——定性认识")
    for para_text in data.get('qualitative_analysis', []):
        add_body(doc, para_text)


def build_section5(doc, data):
    """第五部分：当下定价与交易逻辑"""
    add_h1(doc, "第五部分：思考第二阶段——当下定价与交易逻辑")

    s5 = data.get('trading_logic', {})

    sub_sections = [
        ("5.1 当前市场叙事",       'current_narrative'),
        ("5.2 当前定价隐含的盈利预期", 'implied_earnings'),
        ("5.3 主要分歧",           'divergence'),
    ]
    for title, key in sub_sections:
        if s5.get(key):
            add_h2(doc, title)
            for para_text in s5[key]:
                add_body(doc, para_text)

    # 催化剂 / 风险（bullet）
    if s5.get('catalysts'):
        add_h2(doc, "5.4 上行催化剂")
        for item in s5['catalysts']:
            add_bullet(doc, item)

    if s5.get('risks'):
        add_h2(doc, "5.5 下行风险")
        for item in s5['risks']:
            add_bullet(doc, item)

    # 横向比较表
    if s5.get('peer_comparison'):
        add_h2(doc, "5.6 横向比较")
        peers = s5['peer_comparison']
        headers = ["公司", "核心优势", "预测净利", "成长性", "风险特征"]
        col_widths = [1200, 2800, 1200, 2200, 2320]
        rows = [[p.get('company'), p.get('advantage'), p.get('profit_est'),
                 p.get('growth'), p.get('risk')] for p in peers]
        make_data_table(doc, headers, rows, col_widths)
        doc.add_paragraph()

    if s5.get('recommendation'):
        add_h2(doc, "5.7 交易建议")
        for para_text in s5['recommendation']:
            add_body(doc, para_text)

    # 结尾
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("—— 报告结束 ——")
    set_run_font(r, color=GRAY, size_pt=10)

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    r2 = p2.add_run("免责声明：本报告基于公开信息及用户提供的材料编制，仅供内部参考，不构成任何形式的投资建议。投资有风险，入市需谨慎。")
    set_run_font(r2, color=GRAY, size_pt=9)


# ═══════════════════════════════════════════════
# 主流程
# ═══════════════════════════════════════════════

def generate_report(data: dict, output_path: str):
    doc = Document()
    setup_page(doc)

    meta = data['meta']

    # 页眉 / 页脚
    section = doc.sections[0]
    build_header(section, meta['company_name_cn'], meta['ticker'], meta['report_date'])
    build_footer(section)

    # 封面
    build_cover(doc, meta)

    # 五个部分
    build_section1(doc, data)
    build_section2(doc, data)
    build_section3(doc, data)
    build_section4(doc, data)
    build_section5(doc, data)

    doc.save(output_path)
    _patch_zoom(output_path)
    print(f"✓ 报告已生成：{output_path}")


def _patch_zoom(docx_path):
    """修复 python-docx 生成的 settings.xml 中 zoom 缺少 percent 属性的问题"""
    import zipfile, shutil, tempfile, os
    tmp = docx_path + ".tmp"
    with zipfile.ZipFile(docx_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/settings.xml':
                    data = data.replace(
                        b'<w:zoom w:val=',
                        b'<w:zoom w:percent="100" w:val='
                    )
                zout.writestr(item, data)
    os.replace(tmp, docx_path)


def main():
    parser = argparse.ArgumentParser(description='生成深度研究报告 Word 文档')
    parser.add_argument('--data',   required=True, help='JSON 数据文件路径')
    parser.add_argument('--output', required=True, help='输出 .docx 文件路径')
    args = parser.parse_args()

    with open(args.data, encoding='utf-8') as f:
        data = json.load(f)

    generate_report(data, args.output)


if __name__ == '__main__':
    main()
